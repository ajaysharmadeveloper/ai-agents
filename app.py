import streamlit as st
import os
import tempfile
from typing import List, Dict, Any
import PyPDF2
import docx
from io import BytesIO
import hashlib

# LangChain imports - FIXED for OpenAI version compatibility
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.schema import Document

# Updated OpenAI imports to fix compatibility issues
try:
    from langchain_openai import OpenAIEmbeddings, ChatOpenAI
    from langchain_community.callbacks.manager import get_openai_callback
except ImportError:
    # Fallback for older versions
    from langchain.embeddings.openai import OpenAIEmbeddings
    from langchain.chat_models import ChatOpenAI
    from langchain.callbacks import get_openai_callback

# Set page config
st.set_page_config(
    page_title="Document Intelligence Playground",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
<style>
.main-header {
    font-size: 2.5rem;
    font-weight: bold;
    color: #1f77b4;
    text-align: center;
    margin-bottom: 2rem;
}
.sub-header {
    font-size: 1.2rem;
    color: #666;
    text-align: center;
    margin-bottom: 2rem;
}
.upload-section {
    background-color: #f0f2f6;
    padding: 2rem;
    border-radius: 10px;
    margin-bottom: 2rem;
}
.chat-container {
    background-color: #ffffff;
    padding: 1rem;
    border-radius: 10px;
    border: 1px solid #e0e0e0;
    margin-bottom: 1rem;
}
.user-message {
    background-color: #e3f2fd;
    padding: 1rem;
    border-radius: 10px;
    margin-bottom: 1rem;
}
.assistant-message {
    background-color: #f5f5f5;
    padding: 1rem;
    border-radius: 10px;
    margin-bottom: 1rem;
}
.tech-stack {
    font-size: 0.9rem;
    color: #888;
    text-align: center;
    margin-top: 2rem;
}
</style>
""", unsafe_allow_html=True)


def initialize_session_state():
    """Initialize session state variables"""
    if 'vectorstore' not in st.session_state:
        st.session_state.vectorstore = None
    if 'conversation' not in st.session_state:
        st.session_state.conversation = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'document_processed' not in st.session_state:
        st.session_state.document_processed = False
    if 'document_name' not in st.session_state:
        st.session_state.document_name = None
    if 'processing_stats' not in st.session_state:
        st.session_state.processing_stats = {}


def extract_text_from_pdf(pdf_file) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""


def extract_text_from_docx(docx_file) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""


def extract_text_from_txt(txt_file) -> str:
    """Extract text from TXT file"""
    try:
        return txt_file.read().decode('utf-8')
    except Exception as e:
        st.error(f"Error reading TXT: {str(e)}")
        return ""


def process_document(uploaded_file) -> List[Document]:
    """Process uploaded document and return text chunks"""

    # Extract text based on file type
    if uploaded_file.type == "application/pdf":
        text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = extract_text_from_docx(uploaded_file)
    elif uploaded_file.type == "text/plain":
        text = extract_text_from_txt(uploaded_file)
    else:
        st.error("Unsupported file type")
        return []

    if not text.strip():
        st.error("No text could be extracted from the document")
        return []

    # Split text into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", " ", ""]
    )

    chunks = text_splitter.split_text(text)

    # Create Document objects
    documents = [Document(page_content=chunk) for chunk in chunks]

    # Store processing stats
    st.session_state.processing_stats = {
        'total_characters': len(text),
        'total_chunks': len(documents),
        'avg_chunk_size': len(text) // len(documents) if documents else 0
    }

    return documents


def create_vectorstore(documents: List[Document]):
    """Create vector store from documents"""
    try:
        # Initialize embeddings with explicit API key
        embeddings = OpenAIEmbeddings(
            openai_api_key=os.environ.get("OPENAI_API_KEY")
        )

        # Create vector store
        vectorstore = FAISS.from_documents(documents, embeddings)

        return vectorstore
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")
        st.error("Please ensure you have the latest versions of langchain and openai installed.")
        st.code("pip install --upgrade langchain-openai langchain-community openai")
        return None


def create_conversation_chain(vectorstore):
    """Create conversation chain with memory"""
    try:
        # Initialize LLM - Using ChatOpenAI instead of OpenAI for better compatibility
        llm = ChatOpenAI(
            temperature=0.7,
            model_name="gpt-3.5-turbo",
            openai_api_key=os.environ.get("OPENAI_API_KEY")
        )

        # Initialize memory
        memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="answer"
        )

        # Create conversation chain
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=vectorstore.as_retriever(search_kwargs={"k": 5}),
            memory=memory,
            return_source_documents=True,
            verbose=True
        )

        return conversation_chain
    except Exception as e:
        st.error(f"Error creating conversation chain: {str(e)}")
        return None


def handle_user_input(user_question: str):
    """Handle user input and generate response"""
    if st.session_state.conversation is None:
        st.error("Please upload and process a document first!")
        return

    try:
        with get_openai_callback() as cb:
            response = st.session_state.conversation({'question': user_question})

            # Add to chat history
            st.session_state.chat_history.append({
                'question': user_question,
                'answer': response['answer'],
                'source_documents': response.get('source_documents', []),
                'tokens_used': cb.total_tokens if hasattr(cb, 'total_tokens') else 0,
                'cost': cb.total_cost if hasattr(cb, 'total_cost') else 0.0
            })

    except Exception as e:
        st.error(f"Error processing question: {str(e)}")


def display_chat_history():
    """Display chat history in chatbot style"""
    for i, chat in enumerate(st.session_state.chat_history):
        # User message
        st.markdown(f"""
        <div style="display: flex; justify-content: flex-end; margin-bottom: 10px;">
            <div style="background-color: #007bff; color: white; padding: 10px 15px; border-radius: 18px; max-width: 70%; word-wrap: break-word;">
                {chat['question']}
            </div>
        </div>
        """, unsafe_allow_html=True)

        # Assistant message
        st.markdown(f"""
        <div style="display: flex; justify-content: flex-start; margin-bottom: 20px;">
            <div style="background-color: #f1f1f1; color: #333; padding: 10px 15px; border-radius: 18px; max-width: 70%; word-wrap: break-word;">
                {chat['answer']}
            </div>
        </div>
        """, unsafe_allow_html=True)

        # Show source documents in expander (smaller, less prominent)
        if chat['source_documents']:
            with st.expander(f"📄 View Sources ({len(chat['source_documents'])} references)", expanded=False):
                for j, doc in enumerate(chat['source_documents']):
                    st.markdown(f"**Reference {j + 1}:**")
                    st.text(doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content)

        # Show usage stats (smaller)
        if 'tokens_used' in chat and chat['tokens_used'] > 0:
            st.caption(f"💰 Tokens: {chat['tokens_used']} | Cost: ${chat['cost']:.4f}")


def main():
    """Main application function"""
    initialize_session_state()

    # Header
    st.markdown('<div class="main-header">🧠 Document Intelligence Playground</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Upload any document and ask questions about it using AI</div>',
                unsafe_allow_html=True)

    # Sidebar
    with st.sidebar:
        st.header("🔧 Configuration")

        # API Key input
        api_key = st.text_input("OpenAI API Key", type="password", help="Enter your OpenAI API key")
        if api_key:
            os.environ["OPENAI_API_KEY"] = api_key

        st.header("📊 Document Stats")
        if st.session_state.document_processed:
            # Document name with remove button
            col1, col2 = st.columns([3, 1])
            with col1:
                st.success(f"✅ Document: {st.session_state.document_name}")
            with col2:
                if st.button("❌", key="remove_doc", help="Remove document"):
                    # Reset all document-related session state
                    st.session_state.vectorstore = None
                    st.session_state.conversation = None
                    st.session_state.document_processed = False
                    st.session_state.document_name = None
                    st.session_state.processing_stats = {}
                    st.session_state.chat_history = []
                    st.rerun()

            stats = st.session_state.processing_stats
            st.metric("Total Characters", stats.get('total_characters', 0))
            st.metric("Text Chunks", stats.get('total_chunks', 0))
            st.metric("Avg Chunk Size", stats.get('avg_chunk_size', 0))
        else:
            st.info("No document processed yet")

        st.header("💬 Chat Stats")
        st.metric("Questions Asked", len(st.session_state.chat_history))

        # Clear chat button
        if st.button("🗑️ Clear Chat History"):
            st.session_state.chat_history = []
            st.rerun()

    # Main content area - Upload section moved to top
    st.subheader("📄 Upload Document")

    # Only show upload interface if no document is processed
    if not st.session_state.document_processed:
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['pdf', 'docx', 'txt'],
            help="Upload PDF, DOCX, or TXT files"
        )

        if uploaded_file is not None:
            st.info(f"File uploaded: {uploaded_file.name}")

            if st.button("🔄 Process Document", type="primary"):
                if not api_key:
                    st.error("Please enter your OpenAI API key in the sidebar")
                else:
                    with st.spinner("Processing document..."):
                        # Process document
                        documents = process_document(uploaded_file)

                        if documents:
                            # Create vector store
                            vectorstore = create_vectorstore(documents)

                            if vectorstore:
                                # Create conversation chain
                                conversation = create_conversation_chain(vectorstore)

                                if conversation:
                                    # Update session state
                                    st.session_state.vectorstore = vectorstore
                                    st.session_state.conversation = conversation
                                    st.session_state.document_processed = True
                                    st.session_state.document_name = uploaded_file.name
                                    st.session_state.chat_history = []  # Clear previous chat

                                    st.success("✅ Document processed successfully!")
                                    st.rerun()
    else:
        # Show message when document is already processed
        st.info("📄 Document is ready for questions! Use the ❌ button in the sidebar to upload a different document.")

    # Questions section
    st.subheader("💬 Ask Questions")

    if st.session_state.document_processed:
        # Display chat history ABOVE the question input (chatbot style)
        if st.session_state.chat_history:
            # Create a container with max height for scrolling
            with st.container():
                display_chat_history()

        # Question input - at the bottom
        st.markdown("---")  # Separator line
        user_question = st.text_input(
            "💬 Type your question here:",
            placeholder="What is this document about?",
            key="user_question"
        )

        if st.button("🚀 Ask Question") and user_question:
            with st.spinner("Thinking..."):
                handle_user_input(user_question)
                st.rerun()

        # Sample questions (moved to bottom)
        with st.expander("💡 Try these sample questions"):
            sample_questions = [
                "What is the main topic of this document?",
                "Can you summarize the key points?",
                "What are the important dates mentioned?",
                "Who are the key people or organizations mentioned?",
                "What actions or recommendations are suggested?"
            ]

            for question in sample_questions:
                if st.button(f"💬 {question}", key=f"sample_{question}"):
                    with st.spinner("Thinking..."):
                        handle_user_input(question)
                        st.rerun()
    else:
        st.info("👆 Please upload and process a document first to start asking questions")

    # Footer
    st.markdown("---")
    st.markdown(
        '<div class="tech-stack">Built with LangChain 🦜🔗 | Streamlit | OpenAI GPT | FAISS Vector Store</div>',
        unsafe_allow_html=True
    )

    # Demo information
    with st.expander("ℹ️ About This Demo"):
        st.markdown("""
        **Document Intelligence Playground** demonstrates advanced AI capabilities for document understanding:

        **Key Features:**
        - **Multi-format Support**: PDF, DOCX, and TXT files
        - **Intelligent Chunking**: Optimal text splitting for better retrieval
        - **Semantic Search**: Vector embeddings for context-aware answers
        - **Conversation Memory**: Maintains context across questions
        - **Source Attribution**: Shows relevant document sections
        - **Usage Tracking**: Token usage and cost monitoring

        **Technical Implementation:**
        - **LangChain**: Orchestrates the AI workflow
        - **OpenAI GPT**: Powers the question-answering
        - **FAISS**: Vector database for fast similarity search
        - **Streamlit**: Interactive web interface

        **Business Value:**
        - Instant document understanding
        - Automated knowledge extraction
        - Improved information accessibility
        - Reduced manual document review time
        """)


if __name__ == "__main__":
    main()